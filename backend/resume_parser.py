import PyPDF2
import docx
import re
from typing import Dict, List, Any
import os
from pdf_processor import PDFProcessor
from docx import Document

class ResumeParser:
    def __init__(self, file_path=None):
        self.file_path = file_path
        self.pdf_processor = PDFProcessor()
        self.text = ""
        self.skills = []
        self.experience = []
        self.education = []
        self.contact_info = {}
        # Define regex patterns for extracting information
        self.skill_pattern = re.compile(r'(?i)(?:skills|expertise|proficiencies)[:.\s]+(.*?)(?=\n\n|\Z)', re.DOTALL)
        self.experience_pattern = re.compile(r'(?i)(?:experience|work history|employment)[:.\s]+(.*?)(?=\n\n|\Z)', re.DOTALL)
        self.education_pattern = re.compile(r'(?i)(?:education|academic background|qualifications)[:.\s]+(.*?)(?=\n\n|\Z)', re.DOTALL)

    def parse(self, content: bytes, filename: str):
        """Parse the resume file content and extract information, using filename to determine type."""
        
        if not filename:
            raise ValueError("Filename must be provided to determine file type")
        if not content:
             raise ValueError("File content cannot be empty")
            
        try:
            # Extract text based on the provided filename
            if filename.endswith('.pdf'):
                # Assuming pdf_processor needs a path, save content temporarily?
                # Or modify pdf_processor to accept bytes?
                # TEMPORARY WORKAROUND: Save bytes to temp file (Not ideal for concurrent requests!)
                # A better solution would be to modify PDFProcessor to accept bytes.
                temp_file_path = f"/tmp/{filename}" # Use /tmp dir if available
                try:
                    with open(temp_file_path, "wb") as f:
                        f.write(content)
                    self.text = self.pdf_processor.extract_text(temp_file_path)
                finally:
                    if os.path.exists(temp_file_path):
                         os.remove(temp_file_path)

            elif filename.endswith('.docx'):
                # python-docx can read from a file-like object (BytesIO)
                import io
                doc = Document(io.BytesIO(content))
                self.text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:
                raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

            # Extract information using self.text
            self._extract_skills()
            self._extract_experience()
            self._extract_education()
            # self._extract_contact_info() # Assuming this also uses self.text

            return {
                "text": self.text,
                "skills": self.skills,
                "experience": self.experience,
                "education": self.education,
                # "contact_info": self.contact_info
            }
        except Exception as e:
            print(f"Error parsing resume content for {filename}: {str(e)}")
            # Log the traceback for detailed debugging
            import traceback
            traceback.print_exc()
            raise

    def _extract_skills(self):
        """Extract skills from resume text"""
        try:
            skills_match = self.skill_pattern.search(self.text)
            if skills_match:
                skills_text = skills_match.group(1)
                # Split by commas, semicolons, or newlines and clean up
                skills = [skill.strip() for skill in re.split(r'[,;\n]', skills_text) if skill.strip()]
                self.skills = skills
            else:
                self.skills = []
        except Exception as e:
            print(f"Error extracting skills: {str(e)}")
            self.skills = []

    def _extract_experience(self):
        """Extract work experience from resume text"""
        try:
            experience_match = self.experience_pattern.search(self.text)
            if experience_match:
                experience_text = experience_match.group(1)
                # Split into individual experiences
                experiences = re.split(r'\n(?=\d{4}|\w+\s+\d{4}|Present|Current)', experience_text)
                parsed_experiences = []
                
                for exp in experiences:
                    if not exp.strip():
                        continue
                    
                    # Try to extract title and company
                    title_match = re.search(r'([^•\n]+)(?:at|@|,)\s*([^•\n]+)', exp)
                    if title_match:
                        title = title_match.group(1).strip()
                        company = title_match.group(2).strip()
                    else:
                        title = exp.split('\n')[0].strip()
                        company = "Unknown"
                    
                    # Try to extract duration
                    duration_match = re.search(r'(\d{4}\s*[-–]\s*(?:Present|\d{4}))', exp)
                    duration = duration_match.group(1) if duration_match else "Unknown"
                    
                    parsed_experiences.append({
                        "title": title,
                        "company": company,
                        "duration": duration,
                        "description": exp.strip()
                    })
                
                self.experience = parsed_experiences
            else:
                self.experience = []
        except Exception as e:
            print(f"Error extracting experience: {str(e)}")
            self.experience = []

    def _extract_education(self):
        """Extract education information from resume text"""
        try:
            education_match = self.education_pattern.search(self.text)
            if education_match:
                education_text = education_match.group(1)
                # Split into individual education entries
                education_entries = re.split(r'\n(?=\d{4}|\w+\s+\d{4}|Present|Current)', education_text)
                parsed_education = []
                
                for edu in education_entries:
                    if not edu.strip():
                        continue
                    
                    # Try to extract degree and institution
                    degree_match = re.search(r'([^•\n]+)(?:at|@|,)\s*([^•\n]+)', edu)
                    if degree_match:
                        degree = degree_match.group(1).strip()
                        institution = degree_match.group(2).strip()
                    else:
                        degree = edu.split('\n')[0].strip()
                        institution = "Unknown"
                    
                    # Try to extract year
                    year_match = re.search(r'(\d{4})', edu)
                    year = year_match.group(1) if year_match else "Unknown"
                    
                    parsed_education.append({
                        "degree": degree,
                        "institution": institution,
                        "year": year
                    })
                
                self.education = parsed_education
            else:
                self.education = []
        except Exception as e:
            print(f"Error extracting education: {str(e)}")
            self.education = []

    def _extract_contact_info(self):
        """Extract contact information from resume text"""
        # ... existing code ...

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file"""
        try:
            return self.pdf_processor.extract_text(file_path)
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        try:
            skills_match = self.skill_pattern.search(text)
            if skills_match:
                skills_text = skills_match.group(1)
                # Split by commas, semicolons, or newlines and clean up
                skills = [skill.strip() for skill in re.split(r'[,;\n]', skills_text) if skill.strip()]
                return skills
            return []
        except Exception as e:
            print(f"Error extracting skills: {str(e)}")
            return []

    def extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience from resume text"""
        try:
            experience_match = self.experience_pattern.search(text)
            if experience_match:
                experience_text = experience_match.group(1)
                # Split into individual experiences
                experiences = re.split(r'\n(?=\d{4}|\w+\s+\d{4}|Present|Current)', experience_text)
                parsed_experiences = []
                
                for exp in experiences:
                    if not exp.strip():
                        continue
                    
                    # Try to extract title and company
                    title_match = re.search(r'([^•\n]+)(?:at|@|,)\s*([^•\n]+)', exp)
                    if title_match:
                        title = title_match.group(1).strip()
                        company = title_match.group(2).strip()
                    else:
                        title = exp.split('\n')[0].strip()
                        company = "Unknown"
                    
                    # Try to extract duration
                    duration_match = re.search(r'(\d{4}\s*[-–]\s*(?:Present|\d{4}))', exp)
                    duration = duration_match.group(1) if duration_match else "Unknown"
                    
                    parsed_experiences.append({
                        "title": title,
                        "company": company,
                        "duration": duration,
                        "description": exp.strip()
                    })
                
                return parsed_experiences
            return []
        except Exception as e:
            print(f"Error extracting experience: {str(e)}")
            return []

    def extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education from resume text"""
        try:
            education_match = self.education_pattern.search(text)
            if education_match:
                education_text = education_match.group(1)
                # Split into individual education entries
                education_entries = re.split(r'\n(?=\d{4}|\w+\s+\d{4}|Present|Current)', education_text)
                parsed_education = []
                
                for edu in education_entries:
                    if not edu.strip():
                        continue
                    
                    # Try to extract degree and institution
                    degree_match = re.search(r'([^•\n]+)(?:at|@|,)\s*([^•\n]+)', edu)
                    if degree_match:
                        degree = degree_match.group(1).strip()
                        institution = degree_match.group(2).strip()
                    else:
                        degree = edu.split('\n')[0].strip()
                        institution = "Unknown"
                    
                    # Try to extract year
                    year_match = re.search(r'(\d{4})', edu)
                    year = year_match.group(1) if year_match else "Unknown"
                    
                    parsed_education.append({
                        "degree": degree,
                        "institution": institution,
                        "year": year
                    })
                
                return parsed_education
            return []
        except Exception as e:
            print(f"Error extracting education: {str(e)}")
            return []

    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse a resume file and extract relevant information"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            # Extract text based on file type
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Clean the text
            text = self.pdf_processor.clean_text(text)

            # Extract information
            skills = self.extract_skills(text)
            experience = self.extract_experience(text)
            education = self.extract_education(text)

            return {
                "text": text,
                "skills": skills,
                "experience": experience,
                "education": education
            }
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}") 